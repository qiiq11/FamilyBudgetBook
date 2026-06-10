import os
import sys

# Try installing python-docx to user site
os.system(f'"{sys.executable}" -m pip install python-docx --user --break-system-packages --quiet')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ---- Style helpers ----
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

def add_heading_styled(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.7)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.bold = bold
    return p

def set_cell_font(cell, text, bold=False, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.bold = bold
    return run

def shade_row(row, color="D9E2F3"):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd', {
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill': color,
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'clear'
        })
        shading.append(shading_elm)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_font(hdr.cells[i], h, bold=True, size=10)
    shade_row(hdr, "4472C4")
    for cell in hdr.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            set_cell_font(row.cells[c_idx], str(val), size=10)
        if r_idx % 2 == 0:
            shade_row(row, "D9E2F3")
    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing
    return table

# ===========================================================
# DOCUMENT BODY
# ===========================================================

# ---- Cover / Title ----
title = doc.add_heading('家庭记账本（FamilyBudgetBook）', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '黑体'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('需求分析文档')
run.font.name = '黑体'
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'版本 1.0　|　{datetime.date.today().strftime("%Y年%m月%d日")}')
run.font.name = '宋体'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ===========================================================
# 一、系统概述
# ===========================================================
add_heading_styled('一、系统概述', 1)
add_para(
    '家庭记账本（FamilyBudgetBook）是一个面向家庭用户的小型财务管理 Web 应用，'
    '旨在帮助家庭成员共同管理日常收支账目，提供记账、查询、统计与可视化分析等功能。'
    '系统采用 B/S 架构，用户通过浏览器即可访问，支持多用户独立注册和使用，'
    '每个用户拥有自己独立的家庭账本空间。',
    indent=True
)
add_para(
    '系统的核心价值在于：让家庭成员能够便捷地记录每一笔收入和支出，'
    '清楚地了解"钱从哪里来、花到哪里去"，并通过丰富的统计图表辅助家庭财务决策。',
    indent=True
)

# ===========================================================
# 二、用户角色分析
# ===========================================================
add_heading_styled('二、用户角色分析', 1)

add_para(
    '本系统面向家庭场景，用户群体比较单一清晰，不涉及复杂的角色权限体系。'
    '以下从"系统身份"和"使用视角"两个维度进行分析。',
    indent=True
)

add_heading_styled('2.1 按系统身份划分', 2)

add_table(
    ['角色', '说明'],
    [
        ['未注册用户（访客）', '只能访问登录/注册页面，无法使用任何记账功能。'],
        ['已注册用户（家庭账户）',
         '系统的核心使用者。注册后拥有独立的家庭账本空间（成员、分类、账目均隔离），'
         '可使用系统的全部功能。'],
    ],
    col_widths=[4, 12]
)

add_para(
    '说明：本系统不设"管理员"角色——每个注册用户就是自己家庭账本的管理员，'
    '对自己账本内的数据有完全的增删改查权限。同一家庭中的多个成员（如爸爸、妈妈、孩子）'
    '共用同一个注册账号，通过在"家庭成员"模块中维护成员列表来实现账目归属。',
    indent=True
)

add_heading_styled('2.2 按使用视角划分', 2)

add_para('从家庭记账的实际使用场景出发，可以将使用者理解为以下三类"视角角色"：', indent=True)

add_para('（1）家庭账本管理者', bold=True)
add_para(
    '通常为家庭中的一人（如父母一方），是系统的主要操作者。负责：注册账号、'
    '维护家庭成员信息、设置收支分类、录入家庭日常收支、查看统计报表、管理所有账目数据。',
    indent=True
)

add_para('（2）家庭成员（被记账对象）', bold=True)
add_para(
    '家庭成员本身不直接登录系统，但每一笔收支都会关联到具体的家庭成员，'
    '以便统计每个人相关的收支情况。典型成员如：爸爸、妈妈、孩子、爷爷、奶奶等。',
    indent=True
)

add_para('（3）家庭其他查看者', bold=True)
add_para(
    '同一家庭中的其他成员（如夫妻双方）可以共用同一账号查看家庭财务状况，'
    '了解收支汇总、分类分布和趋势变化。',
    indent=True
)

# ===========================================================
# 三、用户能"用系统做什么"（用例描述）
# ===========================================================
add_heading_styled('三、用户能"用系统做什么"（用例描述）', 1)

add_heading_styled('3.1 未注册用户', 2)
add_para('• 注册账号：填写用户名、密码、家庭名称（可选），注册成功后系统自动创建默认的示例成员'
         '（爸爸、妈妈、小明）和常用收支分类（工资、奖金、餐饮、交通等 12 个），'
         '用户可在此基础上快速开始使用。', indent=True)

add_heading_styled('3.2 已注册用户', 2)

add_para('【身份认证方面】', bold=True)
add_para('• 使用用户名和密码登录系统', indent=True)
add_para('• 查看当前登录账号信息（家庭名称）', indent=True)
add_para('• 退出登录', indent=True)

add_para('【家庭成员维护方面】', bold=True)
add_para('• 添加家庭成员（姓名 + 关系，如"爸爸-父亲""小明-孩子"）', indent=True)
add_para('• 修改已添加的成员信息', indent=True)
add_para('• 删除不再需要的成员（只有该成员下没有关联账目时才能删除，防止数据丢失）', indent=True)

add_para('【收支分类维护方面】', bold=True)
add_para('• 查看收入分类列表和支出分类列表（分左右两栏展示）', indent=True)
add_para('• 新增自定义收支分类（指定名称、收入/支出类型、排序权重）', indent=True)
add_para('• 编辑已有分类的名称和排序', indent=True)
add_para('• 删除不再需要的分类（与成员一样，有关联账目时不允许删除）', indent=True)

add_para('【日常记账方面】', bold=True)
add_para('• 记录一笔收支：选择类型 → 选择成员 → 选择分类 → 填写金额 → 选择日期 → 填写备注（可选）', indent=True)
add_para('• 编辑之前记录的账目：修改任意字段', indent=True)
add_para('• 删除误录的账目（带确认弹窗，防止误删）', indent=True)

add_para('【查询账目方面】', bold=True)
add_para('• 按收支类型筛选（只看收入 / 只看支出）', indent=True)
add_para('• 按家庭成员筛选（只看某个人的账目）', indent=True)
add_para('• 按分类筛选（只看某一类收支）', indent=True)
add_para('• 按日期范围筛选（自定义起止日期）', indent=True)
add_para('• 按关键词搜索（在备注和分类名称中模糊匹配）', indent=True)
add_para('• 多条件组合筛选（以上条件可同时使用）', indent=True)
add_para('• 分页浏览（每页 15 条，可切换页码）', indent=True)

add_para('【首页概览方面】', bold=True)
add_para('• 查看本月收入、支出、结余三个关键数字卡片', indent=True)
add_para('• 查看当前统计周期的文字摘要说明', indent=True)
add_para('• 查看支出分类占比饼图（直观了解钱花在哪里）', indent=True)
add_para('• 查看近 30 天收支趋势折线图', indent=True)
add_para('• 查看最近 8 条账目记录列表', indent=True)

add_para('【统计分析方面】', bold=True)
add_para('• 按不同周期汇总：日、周、月、年，或自定义起止日期', indent=True)
add_para('• 按成员筛选统计数据（只看某个人的收支汇总）', indent=True)
add_para('• 查看文字摘要：区间 + 收入合计及笔数 + 支出合计及笔数 + 结余', indent=True)
add_para('• 查看按成员统计表格：每个成员的收入、支出、结余', indent=True)
add_para('• 查看按分类统计表格：支出分类明细（金额 + 笔数）、收入分类明细', indent=True)
add_para('• 图形化展示：支出分类饼图、收入分类饼图、成员收支对比柱状图、'
         '收支结余趋势组合图（折线+柱状）', indent=True)

# ===========================================================
# 四、功能模块划分与功能点
# ===========================================================
add_heading_styled('四、功能模块划分与功能点', 1)

add_para(
    '基于上述用户角色分析和用例描述，系统划分为以下六大功能模块，'
    '各模块的功能点详细说明如下。',
    indent=True
)

# 4.1
add_heading_styled('4.1 用户认证模块', 2)
add_table(
    ['编号', '功能点', '描述'],
    [
        ['A01', '用户注册',
         '输入用户名、密码（至少6位）、家庭名称（可选）；注册成功后自动初始化'
         '默认成员（爸爸、妈妈、小明）和默认收支分类（12个），同时返回 JWT Token 自动登录'],
        ['A02', '用户登录',
         '输入用户名和密码，验证通过后返回 JWT Token 和用户基本信息'],
        ['A03', '获取当前用户信息',
         '通过 Token 获取当前登录用户的 ID、用户名、家庭名称'],
        ['A04', '登录状态保持',
         '前端通过 Pinia 存储 Token 和用户信息，路由守卫拦截未登录访问'],
        ['A05', '退出登录',
         '清除本地 Token 和用户信息，跳转至登录页'],
    ],
    col_widths=[1.5, 2.5, 12]
)

# 4.2
add_heading_styled('4.2 家庭成员管理模块', 2)
add_table(
    ['编号', '功能点', '描述'],
    [
        ['B01', '成员列表',
         '展示当前家庭所有成员（姓名、关系），以表格形式呈现'],
        ['B02', '添加成员',
         '弹窗表单输入姓名（必填）和关系（选填，如"父亲""母亲""孩子"）'],
        ['B03', '编辑成员',
         '弹窗表单修改已有成员的姓名和关系'],
        ['B04', '删除成员',
         '删除前弹出确认对话框，提示"有关联账目时无法删除"；'
         '后端校验：若该成员存在关联的收支记录，返回错误不允许删除'],
    ],
    col_widths=[1.5, 2.5, 12]
)

# 4.3
add_heading_styled('4.3 收支分类管理模块', 2)
add_table(
    ['编号', '功能点', '描述'],
    [
        ['C01', '分类列表',
         '左右分栏展示收入分类和支出分类，包含名称和排序号'],
        ['C02', '新增分类',
         '弹窗表单输入分类名称（必填）、排序号；选择收入或支出类型'],
        ['C03', '编辑分类',
         '弹窗表单修改已有分类的名称和排序号'],
        ['C04', '删除分类',
         '删除前确认，后端校验：若分类有关联账目则不允许删除'
         '（数据库 FOREIGN KEY 约束 ON DELETE RESTRICT）'],
        ['C05', '注册自动初始化',
         '新用户注册时自动创建 4 个收入分类（工资、奖金、兼职、其他收入）'
         '和 8 个支出分类（餐饮、交通、购物、住房、医疗、教育、娱乐、其他支出）'],
    ],
    col_widths=[1.5, 2.5, 12]
)

# 4.4
add_heading_styled('4.4 收支记录管理模块', 2)
add_table(
    ['编号', '功能点', '描述'],
    [
        ['D01', '记一笔（新增）',
         '弹出表单，选择类型（收入/支出）、成员、分类（根据类型动态过滤）、'
         '金额、日期、备注（选填）；保存后刷新列表'],
        ['D02', '编辑记录',
         '点击编辑按钮，弹出预填好原数据的表单，修改后保存'],
        ['D03', '删除记录',
         '点击删除按钮，确认后删除，刷新列表'],
        ['D04', '多条件组合查询',
         '支持按类型、成员、分类、日期范围、关键词'
         '（模糊匹配备注和分类名）组合筛选'],
        ['D05', '分页',
         '每页 15 条，显示总条数，支持翻页'],
        ['D06', '列表展示',
         '表格展示日期、成员、分类、类型标签（绿色"收入"/红色"支出"）、'
         '金额（带正负号）、备注；收入和支出一条列表混合展示'],
    ],
    col_widths=[1.5, 2.5, 12]
)

# 4.5
add_heading_styled('4.5 首页仪表盘模块', 2)
add_table(
    ['编号', '功能点', '描述'],
    [
        ['E01', '月度概览卡片',
         '四张统计卡片：本月收入、本月支出、本月结余'
         '（正值绿色/负值红色）、统计周期标签'],
        ['E02', '文字摘要',
         '以提示框形式展示当前统计区间的完整文字描述'
         '（统计区间 + 收入合计及笔数 + 支出合计及笔数 + 结余）'],
        ['E03', '支出分类占比饼图',
         'ECharts 环形饼图，展示当月各支出分类的金额占比'],
        ['E04', '收支趋势折线图',
         'ECharts 双线图，展示近 30 天的每日收入和支出变化趋势'],
        ['E05', '最近账目列表',
         '表格展示最近 8 条记账记录'
         '（日期、成员、分类、类型标签、金额、备注）'],
    ],
    col_widths=[1.5, 2.5, 12]
)

# 4.6
add_heading_styled('4.6 汇总统计模块', 2)
add_table(
    ['编号', '功能点', '描述'],
    [
        ['F01', '统计周期切换',
         '支持日、周、月、年四个预设周期，以及自定义起止日期范围'],
        ['F02', '参考日期选择',
         '预设周期模式下可切换参考日期（如切换到上个月查看历史数据）'],
        ['F03', '按成员筛选',
         '下拉选择特定家庭成员，统计仅计算该成员的收支'],
        ['F04', '文字摘要',
         '显示统计区间标签 + 收入合计（笔数）+ 支出合计（笔数）+ 结余金额'],
        ['F05', '按成员统计表格',
         '表格展示每位成员的收入、支出、结余'],
        ['F06', '按分类统计表格',
         '分"支出分类明细"和"收入分类明细"两个表格，'
         '展示各分类的合计金额和笔数'],
        ['F07', '支出分类饼图',
         '当前统计区间内各支出分类占比'],
        ['F08', '收入分类饼图',
         '当前统计区间内各收入分类占比'],
        ['F09', '成员收支对比柱状图',
         '每位成员收入与支出并排对比'],
        ['F10', '收支趋势组合图',
         '收入折线 + 支出折线 + 结余柱状图，按日期展示趋势；'
         '当统计周期为"年"时按月份聚合'],
    ],
    col_widths=[1.5, 2.5, 12]
)

# ===========================================================
# 五、功能模块关系
# ===========================================================
add_heading_styled('五、功能模块关系', 1)

add_para(
    '各模块之间的依赖和协作关系如下：',
    indent=True
)

add_para('• 家庭成员管理模块和收支分类管理模块是基础数据模块，为收支记录模块提供下拉选项数据源。', indent=True)
add_para('• 收支记录管理模块是核心业务模块，所有记账操作（增删改查）在此完成。', indent=True)
add_para('• 汇总统计模块和首页仪表盘模块是数据消费模块，依赖收支记录 + 成员 + 分类数据进行聚合计算和可视化展示。', indent=True)
add_para('• 首页仪表盘模块可以理解为汇总统计模块的简化/快捷版，固定以"本月"为统计周期、近 30 天为趋势范围。', indent=True)
add_para('• 用户认证模块是所有功能模块的前置条件，所有业务操作均需通过登录验证。', indent=True)

add_para('')
add_para('关系示意图：', bold=True)

# Build module relationship diagram using a table
add_table(
    ['层次', '模块', '依赖关系'],
    [
        ['入口层', '用户认证（注册/登录）', '所有模块的前置条件'],
        ['基础数据层', '家庭成员管理、收支分类管理', '为收支记录提供数据源'],
        ['核心业务层', '收支记录管理', '依赖基础数据模块，向统计模块提供原始数据'],
        ['数据消费层', '汇总统计、首页仪表盘', '依赖收支记录 + 成员 + 分类数据，进行聚合与可视化'],
    ],
    col_widths=[3, 5, 8]
)

# ===========================================================
# 六、非功能性需求
# ===========================================================
add_heading_styled('六、非功能性需求', 1)

add_table(
    ['需求项', '说明'],
    [
        ['多用户数据隔离',
         '每个注册用户拥有独立的账本空间，成员、分类、账目数据按 user_id 隔离，'
         '用户之间互不可见。'],
        ['界面设计',
         '基于 Element Plus 组件库，采用左侧导航 + 右侧内容的经典后台管理布局，'
         '界面整洁、操作直观。'],
        ['数据安全',
         '密码使用 bcrypt 加盐哈希存储，不以明文形式落盘；用户认证使用 JWT Token 机制；'
         '分类和成员的删除操作受数据库外键约束保护（ON DELETE RESTRICT / CASCADE），'
         '防止误删导致数据不一致。'],
        ['云端部署',
         '已部署至华为云弹性云服务器 ECS（1vCPU / 1GB / Ubuntu 24.04），'
         '使用 Node.js 22 + PM2 进程管理 + Nginx 反向代理，'
         '可通过公网 IP（http://123.60.7.243）直接访问。'],
        ['浏览器兼容性',
         '支持主流现代浏览器（Chrome、Edge、Firefox 最新版本）。'],
        ['技术栈',
         '前端：Vue 3 + TypeScript + Vite + Element Plus + ECharts + Pinia；'
         '后端：Node.js 22 + Express + SQLite（node:sqlite）；'
         '认证：JWT（jsonwebtoken + bcryptjs）。'],
    ],
    col_widths=[3.5, 12.5]
)

# ===========================================================
# Save
# ===========================================================
output_path = '/sessions/lucid-blissful-johnson/mnt/FamilyBudgetBook/02-需求与设计文档/需求分析文档.docx'
try:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
except:
    pass
doc.save(output_path)
print(f'DOCX saved to: {output_path}')
